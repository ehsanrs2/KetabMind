"""Utilities to export answers to PDF or Word formats."""

from __future__ import annotations

import textwrap
from collections.abc import Iterable, Mapping, Sequence
from io import BytesIO
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

_MARGIN = 72
_WRAP_WIDTH = 90


def _as_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def _normalise_bullets(answer: str) -> list[str]:
    cleaned: list[str] = []
    for line in answer.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        for prefix in ("•", "-", "*"):
            if stripped.startswith(prefix):
                stripped = stripped[len(prefix) :].strip()
                break
        cleaned.append(stripped)
    return cleaned or ([answer.strip()] if answer.strip() else [])


def _ensure_sequence(items: Iterable[Any] | None) -> list[str]:
    if not items:
        return []
    result: list[str] = []
    for item in items:
        text = _as_text(item).strip()
        if text:
            result.append(text)
    return result


def _extract_meta(meta: Any) -> dict[str, Any]:
    if isinstance(meta, Mapping):
        return dict(meta)
    return {}


def _format_metric(value: Any) -> str:
    if isinstance(value, int | float):
        formatted = f"{float(value):.2f}"
        return formatted.rstrip("0").rstrip(".") if "." in formatted else formatted
    return "N/A"


def _wrap_lines(text: str, *, width: int = _WRAP_WIDTH) -> Sequence[str]:
    if not text:
        return ("",)
    wrapped = textwrap.wrap(text, width=width)
    return wrapped or (text,)


def _pdf_answer_section(
    pdf_canvas: canvas.Canvas,
    *,
    question: str,
    bullets: Sequence[str],
    citations: Sequence[str],
    coverage_text: str,
    confidence_text: str,
) -> None:
    text_obj = pdf_canvas.beginText(_MARGIN, letter[1] - _MARGIN)
    text_obj.setFont("Helvetica-Bold", 14)
    for line in _wrap_lines(f"Question: {question}"):
        text_obj.textLine(line)

    text_obj.textLine("")

    text_obj.setFont("Helvetica-Bold", 13)
    text_obj.textLine("Answer:")
    text_obj.setFont("Helvetica", 12)
    if bullets:
        for bullet in bullets:
            wrapped = _wrap_lines(bullet)
            for idx, line in enumerate(wrapped):
                prefix = "• " if idx == 0 else "  "
                text_obj.textLine(f"{prefix}{line}")
            text_obj.textLine("")
    else:
        text_obj.textLine("• (No answer provided)")
        text_obj.textLine("")

    text_obj.setFont("Helvetica-Bold", 13)
    text_obj.textLine("Citations:")
    text_obj.setFont("Helvetica", 12)
    if citations:
        for citation in citations:
            for idx, line in enumerate(_wrap_lines(citation)):
                prefix = "• " if idx == 0 else "  "
                text_obj.textLine(f"{prefix}{line}")
    else:
        text_obj.textLine("• None")

    text_obj.textLine("")
    text_obj.setFont("Helvetica-Bold", 13)
    text_obj.textLine("Meta:")
    text_obj.setFont("Helvetica", 12)
    text_obj.textLine(f"Coverage: {coverage_text}")
    text_obj.textLine(f"Confidence: {confidence_text}")

    pdf_canvas.drawText(text_obj)


def _export_pdf(answer_obj: dict[str, Any]) -> bytes:
    question = _as_text(answer_obj.get("question") or "")
    answer = _as_text(answer_obj.get("answer") or "")
    bullets = _normalise_bullets(answer)
    citations = _ensure_sequence(answer_obj.get("citations"))
    meta = _extract_meta(answer_obj.get("meta"))
    coverage_text = _format_metric(meta.get("coverage"))
    confidence_text = _format_metric(meta.get("confidence"))

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter, pageCompression=0)
    _pdf_answer_section(
        pdf,
        question=question,
        bullets=bullets,
        citations=citations,
        coverage_text=coverage_text,
        confidence_text=confidence_text,
    )
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _export_docx(answer_obj: dict[str, Any]) -> bytes:
    question = _as_text(answer_obj.get("question") or "")
    answer = _as_text(answer_obj.get("answer") or "")
    bullets = _normalise_bullets(answer)
    citations = _ensure_sequence(answer_obj.get("citations"))
    meta = _extract_meta(answer_obj.get("meta"))
    coverage_text = _format_metric(meta.get("coverage"))
    confidence_text = _format_metric(meta.get("confidence"))

    document = Document()

    question_paragraph = document.add_paragraph()
    question_run = question_paragraph.add_run("Question: ")
    question_run.bold = True
    question_paragraph.add_run(question or "—")

    answer_heading = document.add_paragraph()
    answer_heading_run = answer_heading.add_run("Answer:")
    answer_heading_run.bold = True

    if bullets:
        for bullet in bullets:
            document.add_paragraph(bullet, style="List Bullet")
    else:
        document.add_paragraph("No answer provided.", style="List Bullet")

    citations_heading = document.add_paragraph()
    citations_heading_run = citations_heading.add_run("Citations:")
    citations_heading_run.bold = True
    if citations:
        for citation in citations:
            document.add_paragraph(citation, style="List Bullet")
    else:
        document.add_paragraph("None", style="List Bullet")

    meta_heading = document.add_paragraph()
    meta_heading_run = meta_heading.add_run("Meta:")
    meta_heading_run.bold = True
    document.add_paragraph(f"Coverage: {coverage_text}")
    document.add_paragraph(f"Confidence: {confidence_text}")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_answer(answer_obj: dict[str, Any], format: str = "pdf") -> bytes:
    """Export a question/answer payload into the requested document format."""

    fmt = (format or "pdf").strip().lower()
    if fmt == "pdf":
        return _export_pdf(answer_obj)
    if fmt in {"word", "docx"}:
        return _export_docx(answer_obj)
    raise ValueError(f"Unsupported export format: {format}")
